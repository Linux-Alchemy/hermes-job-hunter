"""Bounded career-document rendering and QA for Hermes Job Hunter."""

from __future__ import annotations

import hashlib
import json
import os
import re
import shutil
import subprocess
import tempfile
import unicodedata
import zipfile
from datetime import date, datetime, timezone
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

PLUGIN_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = PLUGIN_DIR / "templates" / "resume_template.docx"
MAX_PAGES = 2
MAX_PDF_BYTES = 2_500_000
MAX_SOURCE_BYTES = 100_000
MAX_SLUG_LENGTH = 100
MAX_VERSION = 9_999
MAX_ANCHOR_LENGTH = 500
REQUIRED_FONT = "JetBrainsMono NF"
REQUIRED_FONT_ALIASES = {REQUIRED_FONT, "JetBrainsMono Nerd Font"}
REQUIRED_PDF_FONT_PREFIXES = ("JetBrainsMonoNF-", "JetBrainsMonoNerdFont-")
MIN_FONT_PT = 9.0
MAX_FONT_PT = 12.0
PAGE_LIMITS = {"resume": 2, "cover_letter": 1}
CONTROLLED_STYLES = {
    "Resume Body",
    "Resume Name",
    "Resume Contact",
    "Resume Section",
    "Resume Entry",
    "Resume Bullet",
    "Cover Letter Header",
    "Cover Letter Subject",
    "Cover Letter Body",
}
def _default_career_root() -> Path:
    """Return the active profile's private workspace root."""
    try:
        from hermes_constants import get_hermes_home

        return (Path(get_hermes_home()) / "workspace").resolve()
    except Exception:
        configured = os.environ.get("HERMES_HOME")
        base = Path(configured).expanduser() if configured else Path.home() / ".hermes"
        return (base / "workspace").resolve()


CAREER_ROOT = _default_career_root()
ALLOWED_LANES = {"applications", "resumes"}
SLUG_RE = re.compile(r"^[a-z0-9]+(?:_[a-z0-9]+)*$")
DATE_RE = re.compile(r"^\d{4}-\d{2}-\d{2}$")
PLACEHOLDER_RE = re.compile(
    r"\{\{|\}\}|\b(?:TODO|TBD|UNKNOWN|VERIFY_BEFORE_USE|PROPOSED)\b",
    re.IGNORECASE,
)
W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _json(payload: dict[str, Any]) -> str:
    return json.dumps(payload, ensure_ascii=False, indent=2, default=str)


def available() -> bool:
    try:
        import docx  # noqa: F401
    except Exception:
        return False
    return TEMPLATE_PATH.is_file() and all(
        shutil.which(command)
        for command in ("soffice", "pdfinfo", "pdftotext", "pdftoppm", "pdffonts", "fc-match")
    )


def _required_font_available() -> tuple[bool, str]:
    matcher = shutil.which("fc-match")
    if not matcher:
        return False, "fontconfig fc-match is unavailable"
    try:
        result = subprocess.run(
            [matcher, "-f", "%{family}", REQUIRED_FONT],
            check=True,
            text=True,
            capture_output=True,
            timeout=10,
        )
    except Exception as exc:
        return False, f"font lookup failed: {exc}"
    families = {item.strip() for item in result.stdout.split(",") if item.strip()}
    matched = sorted(families & REQUIRED_FONT_ALIASES)
    if not matched:
        returned = result.stdout.strip() or "no family"
        return False, f"{REQUIRED_FONT} is not installed (fontconfig returned {returned})"
    return True, ", ".join(matched)


def _fail(message: str, **extra: Any) -> str:
    return _json({"success": False, "error": message, **extra})


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _contained(path: Path, root: Path) -> bool:
    try:
        path.resolve().relative_to(root.resolve())
        return True
    except ValueError:
        return False


def _packet_root(path: Path) -> Path:
    resolved = path.resolve()
    relative = resolved.relative_to(CAREER_ROOT)
    if not relative.parts or relative.parts[0] not in ALLOWED_LANES:
        raise ValueError("Path must be under the active profile workspace applications/ or resumes/")
    if len(relative.parts) < 3:
        raise ValueError("Path is not inside a complete résumé packet")
    if relative.parts[0] == "applications":
        if len(relative.parts) < 4:
            raise ValueError("Application paths require company and role folders")
        return CAREER_ROOT.joinpath(*relative.parts[:3])
    return CAREER_ROOT.joinpath(*relative.parts[:2])


def _validate_source_path(path: Path) -> Path:
    resolved = path.expanduser().resolve(strict=True)
    packet = _packet_root(resolved)
    source_dir = packet / "source"
    if not _contained(resolved, source_dir):
        raise ValueError(f"Approved Markdown must be under {source_dir}")
    if resolved.suffix.lower() != ".md":
        raise ValueError("Approved source must be a Markdown file")
    return packet


def _validate_artifact_path(path: Path, packet: Path, suffix: str) -> Path:
    resolved = path.expanduser().resolve(strict=True)
    if not _contained(resolved, packet / "outputs"):
        raise ValueError(f"Artifact must be under {packet / 'outputs'}")
    if resolved.suffix.lower() != suffix:
        raise ValueError(f"Expected a {suffix} artifact")
    return resolved


def _validated_output_dir(packet: Path) -> Path:
    output_dir = packet / "outputs"
    if output_dir.exists():
        if output_dir.is_symlink() or not output_dir.is_dir():
            raise ValueError("Packet outputs path must be a real directory, not a symlink")
    else:
        output_dir.mkdir(mode=0o700, parents=True)
    if not _contained(output_dir.resolve(), packet.resolve()):
        raise ValueError("Packet outputs path escapes the packet root")
    return output_dir


def _split_frontmatter(text: str) -> tuple[dict[str, Any], str]:
    if not text.startswith("---\n"):
        return {}, text
    end = text.find("\n---\n", 4)
    if end < 0:
        return {}, text
    import yaml

    loaded = yaml.load(text[4:end], Loader=yaml.BaseLoader)
    if not isinstance(loaded, dict):
        raise ValueError("Frontmatter must be a YAML mapping")
    fields = {str(key): value for key, value in loaded.items()}
    return fields, text[end + 5 :]


def _approval_check(source: Path) -> tuple[dict[str, Any], str]:
    if source.stat().st_size > MAX_SOURCE_BYTES:
        raise ValueError(f"Approved source exceeds {MAX_SOURCE_BYTES} bytes")
    fields, body = _split_frontmatter(source.read_text(encoding="utf-8"))
    errors = []
    document_type = fields.get("document_type", "")
    if document_type not in PAGE_LIMITS:
        errors.append("document_type must be resume or cover_letter")
    if fields.get("lifecycle_state") != "APPROVED_MARKDOWN":
        errors.append("lifecycle_state must be APPROVED_MARKDOWN")
    approver = str(fields.get("approved_by") or "").strip()
    if approver.casefold() in {"", "null", "none", "unknown"}:
        errors.append("approved_by must name the human approver")
    approved_at = fields.get("approved_at")
    try:
        if isinstance(approved_at, datetime):
            raise ValueError
        if isinstance(approved_at, date):
            approved_date = approved_at
        else:
            approved_date = date.fromisoformat(str(approved_at))
        if approved_date.isoformat() != str(approved_at):
            raise ValueError
    except (TypeError, ValueError):
        errors.append("approved_at must be YYYY-MM-DD")
    if errors:
        raise ValueError("; ".join(errors))
    if PLACEHOLDER_RE.search(body):
        raise ValueError("Approved source still contains a placeholder or unresolved evidence label")
    if document_type == "resume" and not re.search(r"^#\s+\S", body, re.MULTILINE):
        raise ValueError("Approved source needs a level-1 name/title heading")
    if document_type == "cover_letter" and not re.search(r"^Dear\s+\S", body, re.MULTILINE):
        raise ValueError("Approved cover letter needs a salutation beginning with Dear")
    return fields, body


def _plain_markdown(body: str) -> str:
    lines = []
    for raw in body.splitlines():
        line = raw.strip()
        if not line:
            continue
        line = re.sub(r"^#{1,6}\s+", "", line)
        line = re.sub(r"^[-*+]\s+", "", line)
        line = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
        line = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", line)
        line = re.sub(r"`([^`]+)`", r"\1", line)
        lines.append(line)
    return "\n".join(lines)


def _tokens(text: str) -> list[str]:
    normal = unicodedata.normalize("NFKC", text).casefold()
    return re.findall(r"[\w@.+:/-]+", normal, flags=re.UNICODE)


def _add_inline(paragraph: Any, text: str) -> None:
    pattern = re.compile(r"(\*\*.+?\*\*|(?<!\*)\*[^*]+?\*(?!\*)|`[^`]+?`)")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            paragraph.add_run(token[1:-1])
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def _render_resume(doc: Any, body: str) -> None:
    before_first_section = True
    for raw in body.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(style="Resume Name")
            _add_inline(p, line[2:].strip())
        elif line.startswith("## "):
            before_first_section = False
            p = doc.add_paragraph(style="Resume Section")
            _add_inline(p, line[3:].strip().upper())
        elif line.startswith("### "):
            p = doc.add_paragraph(style="Resume Entry")
            _add_inline(p, line[4:].strip())
        elif re.match(r"^[-*+]\s+", line):
            p = doc.add_paragraph(style="Resume Bullet")
            _add_inline(p, "• " + re.sub(r"^[-*+]\s+", "", line))
        else:
            style = "Resume Contact" if before_first_section else "Resume Body"
            p = doc.add_paragraph(style=style)
            _add_inline(p, line)


def _render_cover_letter(doc: Any, body: str) -> None:
    blocks = [block for block in re.split(r"\n\s*\n", body.strip()) if block.strip()]
    for index, block in enumerate(blocks):
        lines = [line.strip() for line in block.splitlines() if line.strip()]
        if not lines:
            continue
        if index == 0 or len(lines) > 1:
            for line in lines:
                p = doc.add_paragraph(style="Cover Letter Header")
                _add_inline(p, line)
            doc.paragraphs[-1].paragraph_format.space_after = doc.styles[
                "Cover Letter Body"
            ].paragraph_format.space_after
            continue
        line = lines[0]
        style = (
            "Cover Letter Subject"
            if re.match(r"^\*\*Re:", line, re.IGNORECASE)
            else "Cover Letter Body"
        )
        p = doc.add_paragraph(style=style)
        _add_inline(p, line)


def _render_docx(body: str, output: Path, title: str, document_type: str) -> None:
    from docx import Document

    doc = Document(str(TEMPLATE_PATH))
    for paragraph in list(doc.paragraphs):
        paragraph._element.getparent().remove(paragraph._element)

    if document_type == "resume":
        _render_resume(doc, body)
    elif document_type == "cover_letter":
        _render_cover_letter(doc, body)
    else:
        raise ValueError(f"Unsupported document type: {document_type}")

    doc.core_properties.title = title
    doc.core_properties.subject = "Résumé" if document_type == "resume" else "Cover Letter"
    doc.core_properties.author = ""
    doc.core_properties.keywords = document_type
    doc.save(output)


def _run(argv: list[str], timeout: int = 90) -> subprocess.CompletedProcess[str]:
    return subprocess.run(argv, check=True, text=True, capture_output=True, timeout=timeout)


def _convert_pdf(docx_path: Path, out_dir: Path, user_profile: Path) -> Path:
    user_profile.mkdir(parents=True, exist_ok=True)
    profile_uri = user_profile.resolve().as_uri()
    _run([
        shutil.which("soffice") or "soffice",
        "--headless",
        f"-env:UserInstallation={profile_uri}",
        "--convert-to",
        "pdf",
        "--outdir",
        str(out_dir),
        str(docx_path),
    ])
    pdf = out_dir / f"{docx_path.stem}.pdf"
    if not pdf.is_file():
        raise RuntimeError("LibreOffice did not produce the expected PDF")
    return pdf


def _docx_text_and_shape(path: Path) -> tuple[str, dict[str, Any]]:
    with zipfile.ZipFile(path) as archive:
        names = set(archive.namelist())
        document = ET.fromstring(archive.read("word/document.xml"))
        external_relationships: list[str] = []
        for name in sorted(item for item in names if item.endswith(".rels")):
            relationships = ET.fromstring(archive.read(name))
            for relationship in relationships:
                if relationship.attrib.get("TargetMode") == "External":
                    external_relationships.append(relationship.attrib.get("Target", ""))
        columns = []
        for node in document.iter(f"{W_NS}cols"):
            columns.append(int(node.attrib.get(f"{W_NS}num", "1")))
        text = "\n".join(
            "".join(node.text or "" for node in paragraph.iter(f"{W_NS}t"))
            for paragraph in document.iter(f"{W_NS}p")
        )
        shape = {
            "tables": len(list(document.iter(f"{W_NS}tbl"))),
            "drawings": len(list(document.iter(f"{W_NS}drawing"))),
            "textboxes": len(list(document.iter(f"{W_NS}txbxContent"))),
            "headers": sorted(name for name in names if re.fullmatch(r"word/header\d+\.xml", name)),
            "footers": sorted(name for name in names if re.fullmatch(r"word/footer\d+\.xml", name)),
            "macros": sorted(name for name in names if "vbaProject" in name or name.endswith(".bin")),
            "external_relationships": external_relationships,
            "embedded_objects": sorted(name for name in names if name.startswith("word/embeddings/")),
            "comments": sorted(name for name in names if name.startswith("word/comments")),
            "alt_chunks": len(list(document.iter(f"{W_NS}altChunk"))),
            "tracked_changes": sum(
                len(list(document.iter(f"{W_NS}{tag}")))
                for tag in ("ins", "del", "moveFrom", "moveTo")
            ),
            "column_counts": columns,
        }
    return text, shape


def _docx_font_profile(path: Path) -> dict[str, dict[str, Any]]:
    from docx import Document

    document = Document(str(path))
    profile: dict[str, dict[str, Any]] = {}
    for style in document.styles:
        if style.name not in CONTROLLED_STYLES:
            continue
        size = style.font.size.pt if style.font.size is not None else None
        profile[style.name] = {"family": style.font.name, "size_pt": size}
    return profile


def _pdf_text(path: Path, target: Path) -> str:
    _run([shutil.which("pdftotext") or "pdftotext", "-layout", str(path), str(target)])
    return target.read_text(encoding="utf-8", errors="replace")


def _pdf_pages(path: Path) -> int:
    result = _run([shutil.which("pdfinfo") or "pdfinfo", str(path)])
    match = re.search(r"^Pages:\s+(\d+)$", result.stdout, re.MULTILINE)
    if not match:
        raise RuntimeError("pdfinfo did not report a page count")
    return int(match.group(1))


def _pdf_fonts(path: Path) -> list[str]:
    result = _run([shutil.which("pdffonts") or "pdffonts", str(path)])
    fonts: list[str] = []
    for raw in result.stdout.splitlines():
        line = raw.strip()
        if not line or line.startswith("name") or line.startswith("-"):
            continue
        name = line.split()[0]
        if "+" in name:
            name = name.split("+", 1)[1]
        fonts.append(name)
    if not fonts:
        raise RuntimeError("pdffonts did not report any fonts")
    return sorted(set(fonts))


def _render_previews(path: Path, output_prefix: Path) -> list[Path]:
    _run([
        shutil.which("pdftoppm") or "pdftoppm",
        "-png",
        "-r",
        "150",
        str(path),
        str(output_prefix),
    ])
    return sorted(output_prefix.parent.glob(f"{output_prefix.name}-*.png"))


def _qa(source: Path, docx_path: Path, pdf_path: Path, work_dir: Path) -> dict[str, Any]:
    fields, body = _approval_check(source)
    source_plain = _plain_markdown(body)
    docx_text, shape = _docx_text_and_shape(docx_path)
    font_profile = _docx_font_profile(docx_path)
    extracted = work_dir / f"{pdf_path.stem}_extracted.txt"
    pdf_text = _pdf_text(pdf_path, extracted)
    pages = _pdf_pages(pdf_path)
    pdf_fonts = _pdf_fonts(pdf_path)
    preview_dir = work_dir / f"{pdf_path.stem}_preview"
    preview_dir.mkdir(parents=True, exist_ok=True)
    previews = _render_previews(pdf_path, preview_dir / "page")

    source_tokens = _tokens(source_plain)
    docx_tokens = _tokens(docx_text)
    pdf_tokens = _tokens(pdf_text)
    prefix = "Resume" if fields["document_type"] == "resume" else "Cover"
    expected_styles = {name for name in CONTROLLED_STYLES if name.startswith(prefix)}
    fonts_ok = expected_styles.issubset(font_profile) and all(
        font_profile[name]["family"] == REQUIRED_FONT for name in expected_styles
    )
    font_sizes_ok = expected_styles.issubset(font_profile) and all(
        font_profile[name]["size_pt"] is not None
        and MIN_FONT_PT <= font_profile[name]["size_pt"] <= MAX_FONT_PT
        for name in expected_styles
    )
    page_limit = PAGE_LIMITS[fields["document_type"]]
    checks = {
        "source_is_approved": True,
        "docx_is_valid_zip": zipfile.is_zipfile(docx_path),
        "docx_has_no_tables": shape["tables"] == 0,
        "docx_has_no_drawings": shape["drawings"] == 0,
        "docx_has_no_textboxes": shape["textboxes"] == 0,
        "docx_has_no_headers_or_footers": not shape["headers"] and not shape["footers"],
        "docx_has_no_macros": not shape["macros"],
        "docx_has_no_external_relationships": not shape["external_relationships"],
        "docx_has_no_embedded_objects": not shape["embedded_objects"],
        "docx_has_no_comments": not shape["comments"],
        "docx_has_no_alt_chunks": shape["alt_chunks"] == 0,
        "docx_has_no_tracked_changes": shape["tracked_changes"] == 0,
        "docx_is_single_column": all(count == 1 for count in shape["column_counts"]),
        "docx_text_matches_source": docx_tokens == source_tokens,
        "pdf_text_matches_source": pdf_tokens == source_tokens,
        "pdf_has_text_layer": bool(pdf_text.strip()),
        "docx_uses_required_font": fonts_ok,
        "docx_font_sizes_within_range": font_sizes_ok,
        "pdf_uses_required_fonts": all(
            font.startswith(REQUIRED_PDF_FONT_PREFIXES) for font in pdf_fonts
        ),
        "pdf_page_count_within_limit": 1 <= pages <= page_limit,
        "pdf_below_parser_size_limit": pdf_path.stat().st_size < MAX_PDF_BYTES,
        "preview_count_matches_pages": len(previews) == pages,
        "no_unresolved_placeholders": not PLACEHOLDER_RE.search(source_plain),
    }
    return {
        "validation_passed": all(checks.values()),
        "checks": checks,
        "approval": fields,
        "page_count": pages,
        "page_limit": page_limit,
        "docx_shape": shape,
        "font_standard": {
            "family": REQUIRED_FONT,
            "minimum_pt": MIN_FONT_PT,
            "maximum_pt": MAX_FONT_PT,
        },
        "font_profile": font_profile,
        "pdf_fonts": pdf_fonts,
        "source_sha256": _sha256(source),
        "template": TEMPLATE_PATH.name,
        "template_sha256": _sha256(TEMPLATE_PATH),
        "docx_sha256": _sha256(docx_path),
        "pdf_sha256": _sha256(pdf_path),
        "pdf_bytes": pdf_path.stat().st_size,
        "extracted_text": str(extracted),
        "previews": [str(path) for path in previews],
    }


def _atomic_publish(work_dir: Path, output_dir: Path, stem: str) -> dict[str, Path]:
    names = {
        "docx": f"{stem}.docx",
        "pdf": f"{stem}.pdf",
        "extracted": f"{stem}_extracted.txt",
        "preview": f"{stem}_preview",
        "qa": f"{stem}_qa.json",
    }
    destinations = {key: output_dir / name for key, name in names.items()}
    for path in destinations.values():
        if path.exists():
            raise FileExistsError(f"Refusing to overwrite existing artifact: {path}")
    placed: list[Path] = []
    try:
        for key in ("docx", "pdf", "extracted", "preview", "qa"):
            os.replace(work_dir / names[key], destinations[key])
            placed.append(destinations[key])
    except Exception:
        for path in reversed(placed):
            if path.is_dir():
                shutil.rmtree(path)
            elif path.exists():
                path.unlink()
        raise
    for key in ("docx", "pdf", "extracted", "qa"):
        destinations[key].chmod(0o600)
    return destinations


def _replace_existing_set(work_dir: Path, output_dir: Path, stem: str) -> None:
    names = [
        f"{stem}.docx",
        f"{stem}.pdf",
        f"{stem}_extracted.txt",
        f"{stem}_preview",
        f"{stem}_qa.json",
    ]
    destinations = [output_dir / name for name in names]
    if not all(path.exists() for path in destinations):
        raise FileNotFoundError("Layout revision requires a complete existing artifact set")

    backup = work_dir / "rollback"
    backup.mkdir()
    moved_old: list[tuple[Path, Path]] = []
    moved_new: list[Path] = []
    try:
        for destination in destinations:
            saved = backup / destination.name
            os.replace(destination, saved)
            moved_old.append((saved, destination))
        for destination in destinations:
            os.replace(work_dir / destination.name, destination)
            moved_new.append(destination)
    except Exception:
        for destination in reversed(moved_new):
            if destination.is_dir():
                shutil.rmtree(destination)
            elif destination.exists():
                destination.unlink()
        for saved, destination in reversed(moved_old):
            if saved.exists():
                os.replace(saved, destination)
        raise
    for destination in destinations:
        if destination.is_file():
            destination.chmod(0o600)


def render_approved(args: dict[str, Any], **_: Any) -> str:
    try:
        if not available():
            return _fail("Required renderer dependencies are unavailable")
        font_available, font_detail = _required_font_available()
        if not font_available:
            return _fail(
                f"Required font {REQUIRED_FONT} is unavailable; stop and ask the human decision owner to select an alternate font",
                font_check=font_detail,
            )
        source = Path(str(args.get("source_markdown", "")))
        slug = str(args.get("document_slug", "")).strip()
        if len(slug) > MAX_SLUG_LENGTH or not SLUG_RE.fullmatch(slug):
            return _fail("document_slug must be snake_case")
        version = int(args.get("version", 0))
        if not 1 <= version <= MAX_VERSION:
            return _fail(f"version must be between 1 and {MAX_VERSION}")
        packet = _validate_source_path(source)
        fields, body = _approval_check(source)
        output_dir = _validated_output_dir(packet)
        stem = f"{slug}_v{version}"
        existing = [output_dir / f"{stem}.docx", output_dir / f"{stem}.pdf", output_dir / f"{stem}_qa.json"]
        if any(path.exists() for path in existing):
            return _fail("Refusing to overwrite an existing version", existing=[str(p) for p in existing if p.exists()])

        with tempfile.TemporaryDirectory(prefix=f".{stem}-", dir=output_dir) as temp:
            work = Path(temp)
            docx_path = work / f"{stem}.docx"
            _render_docx(
                body,
                docx_path,
                fields.get("title", slug),
                fields["document_type"],
            )
            pdf_path = _convert_pdf(docx_path, work, work / "lo-profile")
            qa = _qa(source, docx_path, pdf_path, work)
            if not qa["validation_passed"]:
                return _fail(
                    "Rendered artifacts failed QA and were not published",
                    checks=qa["checks"],
                    page_count=qa["page_count"],
                )
            qa.update({
                "tool": "resume_render_approved",
                "generated_at": datetime.now(timezone.utc).isoformat(),
                "source": str(source.resolve().relative_to(packet.resolve())),
                "packet_root": ".",
                "document_slug": slug,
                "version": version,
                "lifecycle_state": "VALIDATED" if qa["validation_passed"] else "RENDERED_VALIDATION_FAILED",
                "ready_for_external_use": False,
                "next_owner": "human decision owner",
                "next_action": "Review rendered page previews and approve or request changes.",
            })
            qa["extracted_text"] = f"{stem}_extracted.txt"
            qa["previews"] = [
                f"{stem}_preview/{Path(path).name}"
                for path in qa["previews"]
            ]
            (work / f"{stem}_qa.json").write_text(_json(qa) + "\n", encoding="utf-8")
            destinations = _atomic_publish(work, output_dir, stem)

        return _json({
            "success": True,
            "validation_passed": qa["validation_passed"],
            "lifecycle_state": qa["lifecycle_state"],
            "artifacts": {key: str(path) for key, path in destinations.items()},
            "checks": qa["checks"],
            "page_count": qa["page_count"],
            "ready_for_external_use": False,
            "next_owner": "human decision owner",
        })
    except Exception as exc:
        return _fail(f"Render failed: {exc}")


def validate_artifacts(args: dict[str, Any], **_: Any) -> str:
    try:
        if not available():
            return _fail("Required validation dependencies are unavailable")
        source = Path(str(args.get("source_markdown", "")))
        packet = _validate_source_path(source)
        output_dir = _validated_output_dir(packet)
        docx_path = _validate_artifact_path(Path(str(args.get("docx_path", ""))), packet, ".docx")
        pdf_path = _validate_artifact_path(Path(str(args.get("pdf_path", ""))), packet, ".pdf")
        with tempfile.TemporaryDirectory(prefix=".resume-qa-", dir=output_dir) as temp:
            qa = _qa(source, docx_path, pdf_path, Path(temp))
            preview_count = len(qa.pop("previews"))
            qa.pop("extracted_text", None)
        return _json({
            "success": True,
            **qa,
            "preview_count": preview_count,
            "preview_artifacts_persisted": False,
            "ready_for_external_use": False,
            "next_owner": "human decision owner",
        })
    except Exception as exc:
        return _fail(f"Validation failed: {exc}")


def _apply_layout_revision(docx_path: Path, operation: str, anchor_text: str) -> dict[str, Any]:
    from docx import Document

    document = Document(str(docx_path))
    matches = [index for index, paragraph in enumerate(document.paragraphs) if paragraph.text.strip() == anchor_text]
    if len(matches) != 1:
        raise ValueError(f"anchor_text must match exactly one paragraph; found {len(matches)}")

    index = matches[0]
    anchor = document.paragraphs[index]
    affected = [anchor.text]
    if operation == "page_break_before":
        anchor.paragraph_format.page_break_before = True
    elif operation == "remove_page_break_before":
        anchor.paragraph_format.page_break_before = False
    elif operation == "keep_block_together":
        block = []
        for paragraph in document.paragraphs[index:]:
            if block and paragraph.style.name in {"Resume Entry", "Resume Section"}:
                break
            block.append(paragraph)
        for position, paragraph in enumerate(block):
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.keep_with_next = position < len(block) - 1
        affected = [paragraph.text for paragraph in block]
    else:
        raise ValueError("Unsupported layout operation")

    document.save(docx_path)
    return {"operation": operation, "anchor_text": anchor_text, "affected_paragraphs": affected}


def revise_layout(args: dict[str, Any], **_: Any) -> str:
    try:
        if not available():
            return _fail("Required renderer dependencies are unavailable")
        font_available, font_detail = _required_font_available()
        if not font_available:
            return _fail(
                f"Required font {REQUIRED_FONT} is unavailable; stop and ask the human decision owner to select an alternate font",
                font_check=font_detail,
            )
        if args.get("human_authorized") is not True:
            return _fail("Layout revision requires explicit human authorization")

        source = Path(str(args.get("source_markdown", "")))
        packet = _validate_source_path(source)
        fields, approved_body = _approval_check(source)
        del approved_body
        original_docx = _validate_artifact_path(Path(str(args.get("docx_path", ""))), packet, ".docx")
        operation = str(args.get("operation", "")).strip()
        anchor_text = str(args.get("anchor_text", "")).strip()
        if not anchor_text or len(anchor_text) > MAX_ANCHOR_LENGTH:
            return _fail(f"anchor_text is required and must not exceed {MAX_ANCHOR_LENGTH} characters")
        if operation == "keep_block_together" and fields["document_type"] != "resume":
            return _fail("keep_block_together is supported only for résumés")

        stem = original_docx.stem
        output_dir = _validated_output_dir(packet)
        original_pdf = output_dir / f"{stem}.pdf"
        original_qa = output_dir / f"{stem}_qa.json"
        if not original_pdf.is_file() or not original_qa.is_file():
            return _fail("Layout revision requires the existing DOCX, PDF, and QA manifest")
        prior_qa = json.loads(original_qa.read_text(encoding="utf-8"))
        if prior_qa.get("ready_for_external_use") is True:
            return _fail("Refusing to edit an artifact already marked ready for external use")

        original_hashes = {"docx": _sha256(original_docx), "pdf": _sha256(original_pdf)}
        with tempfile.TemporaryDirectory(prefix=f".{stem}-layout-", dir=output_dir) as temp:
            work = Path(temp)
            revised_docx = work / f"{stem}.docx"
            shutil.copy2(original_docx, revised_docx)
            revision = _apply_layout_revision(revised_docx, operation, anchor_text)
            revised_pdf = _convert_pdf(revised_docx, work, work / "lo-profile")
            qa = _qa(source, revised_docx, revised_pdf, work)
            if not qa["validation_passed"]:
                return _fail(
                    "Layout revision failed QA; original artifacts were preserved",
                    checks=qa["checks"],
                    page_count=qa["page_count"],
                )
            qa.update({
                "tool": "resume_revise_layout",
                "generated_at": datetime.now(timezone.utc).isoformat(),
                "source": str(source.resolve().relative_to(packet.resolve())),
                "packet_root": ".",
                "lifecycle_state": "VALIDATED" if qa["validation_passed"] else "RENDERED_VALIDATION_FAILED",
                "ready_for_external_use": False,
                "layout_revision": revision,
                "replaced_artifact_hashes": original_hashes,
                "next_owner": "human decision owner",
                "next_action": "Inspect every regenerated page preview, update the visual review, then return the artifact to the human decision owner.",
            })
            qa["extracted_text"] = f"{stem}_extracted.txt"
            qa["previews"] = [f"{stem}_preview/{Path(path).name}" for path in qa["previews"]]
            (work / f"{stem}_qa.json").write_text(_json(qa) + "\n", encoding="utf-8")
            _replace_existing_set(work, output_dir, stem)

        return _json({
            "success": True,
            "validation_passed": qa["validation_passed"],
            "lifecycle_state": qa["lifecycle_state"],
            "layout_revision": revision,
            "artifacts": {
                "docx": str(original_docx),
                "pdf": str(original_pdf),
                "extracted": str(output_dir / f"{stem}_extracted.txt"),
                "preview": str(output_dir / f"{stem}_preview"),
                "qa": str(original_qa),
            },
            "checks": qa["checks"],
            "page_count": qa["page_count"],
            "ready_for_external_use": False,
            "next_owner": "human decision owner",
        })
    except Exception as exc:
        return _fail(f"Layout revision failed: {exc}")
